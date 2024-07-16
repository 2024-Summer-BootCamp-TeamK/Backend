import uuid

import groupdocs_conversion_cloud
from shutil import copyfile
import requests
import os
from docx.shared import Pt
from dotenv import load_dotenv
from docx import Document

from .docxUpload import docx_upload

load_dotenv()


def pdf_convert_docx(url: str, replacement_list: list) -> bytes:
    """
    S3에서 PDF를 다운로드하고, 텍스트 교체 후 Word로 변환하고, 최종적으로 PDF로 변환하여 반환합니다.

    :param url: PDF 파일의 S3 URL
    :param replacement_list: 교체할 문장과 그에 대응하는 대체 문장 튜플
    :return: 수정한 docx파일의 s3 URL
    """

    # PDF 파일의 로컬 경로 설정

    pdf_file = f'{uuid.uuid4()}.pdf'
    docx_file = f'{uuid.uuid4()}.docx'

    corrected_docx_file = 'output_corrected.docx'

    # PDF를 다운로드하여 로컬에 저장
    def download_pdf_from_s3(url: str, local_path: str):
        try:
            response = requests.get(url)
            response.raise_for_status()  # 요청이 성공하지 않을 경우 예외 발생
            with open(local_path, 'wb') as file:
                file.write(response.content)
            print(f"PDF 파일이 {local_path}에 저장되었습니다.")
        except requests.exceptions.RequestException as e:
            raise RuntimeError(f"PDF 다운로드 중 오류 발생: {e}")

    source_url = f'https://{os.getenv("AWS_STORAGE_BUCKET_NAME")}.s3.ap-northeast-2.amazonaws.com/{url}'
    print("\nsource_url: ", source_url)
    download_pdf_from_s3(source_url, pdf_file)

    # PDF 파일이 존재하는지 확인
    if not os.path.isfile(pdf_file):
        raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {pdf_file}")

    try:
        # Get your app_sid and app_key at https://dashboard.groupdocs.cloud (free registration is required).
        app_sid = "c0a01269-583e-4840-a433-acf04f50d795"
        app_key = "80c9eb7a05333a39c977c79aef838574"

        # Create instance of the API
        convert_api = groupdocs_conversion_cloud.ConvertApi.from_keys(app_sid, app_key)
        file_api = groupdocs_conversion_cloud.FileApi.from_keys(app_sid, app_key)

        # upload soruce file to storage
        filename = pdf_file
        remote_name = pdf_file
        output_name = docx_file
        strformat = 'docx'

        request_upload = groupdocs_conversion_cloud.UploadFileRequest(remote_name, filename)
        response_upload = file_api.upload_file(request_upload)

        # Convert PDF to DOCX
        settings = groupdocs_conversion_cloud.ConvertSettings()
        settings.file_path = remote_name
        settings.format = strformat
        settings.output_path = output_name

        loadOptions = groupdocs_conversion_cloud.PdfLoadOptions()
        loadOptions.hide_pdf_annotations = True
        loadOptions.remove_embedded_files = False
        loadOptions.flatten_all_fields = True

        settings.load_options = loadOptions

        convertOptions = groupdocs_conversion_cloud.DocxConvertOptions()
        convertOptions.from_page = 1
        convertOptions.pages_count = 20

        settings.convert_options = convertOptions

        request = groupdocs_conversion_cloud.ConvertDocumentRequest(settings)
        response = convert_api.convert_document(request)
        print("Document converted successfully: " + str(response))

        # Download Document from Storage
        request_download = groupdocs_conversion_cloud.DownloadFileRequest(output_name)
        response_download = file_api.download_file(request_download)
        copyfile(response_download, docx_file)
        print("Result {}".format(response_download))
    except groupdocs_conversion_cloud.ApiException as e:
        print("Exception when calling get_supported_conversion_types:")

    def normalize_text(text):
        # 공백 문자와 줄바꿈 문자를 제거합니다.
        return ''.join(text.split())

    def replace_text_in_docx(doc_path, replacements):
        # Word 문서 열기
        doc = Document(doc_path)

        # 모든 단락의 텍스트를 하나로 결합합니다.
        full_text = '\n'.join([para.text for para in doc.paragraphs])

        # 각 치환 작업을 수행합니다.
        for sub_text, replacement_text in replacements:
            normalized_full_text = normalize_text(full_text)
            normalized_sub_text = normalize_text(sub_text)

            start_idx = normalized_full_text.find(normalized_sub_text)
            while start_idx != -1:
                actual_start_idx = -1
                actual_end_idx = -1
                normalized_index = 0
                main_chars = list(full_text)

                for i, char in enumerate(main_chars):
                    if char not in (' ', '\n'):
                        if normalized_index == start_idx:
                            actual_start_idx = i
                        normalized_index += 1
                    if normalized_index == start_idx + len(normalized_sub_text):
                        actual_end_idx = i + 1
                        break

                # 텍스트를 대체합니다.
                if actual_start_idx != -1 and actual_end_idx != -1:
                    full_text = full_text[:actual_start_idx] + replacement_text + full_text[actual_end_idx:]
                    normalized_full_text = normalize_text(full_text)

                start_idx = normalized_full_text.find(normalized_sub_text)

        # 새로운 Document 객체를 생성하여 수정된 텍스트를 넣습니다.
        new_doc = Document()
        for para_text in full_text.split('\n'):
            paragraph = new_doc.add_paragraph(para_text)

            # 문단 간격 설정 예시
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = Pt(15)  # 20 포인트의 줄 간격 설정

        # 수정된 문서 저장
        corrected_doc_path = f'corrected_{uuid.uuid4()}' + doc_path
        new_doc.save(corrected_doc_path)
        print(f"pdfToDocxWithModify.문서가 성공적으로 저장되었습니다: {corrected_doc_path}")
        return corrected_doc_path

    # 문장 대체 함수 호출
    result = replace_text_in_docx(docx_file, replacement_list)

    uploaded_file_key = docx_upload(result)

    if uploaded_file_key:
        print(f'pdfToDocxWithModify.업로드된 파일의 S3 경로: {uploaded_file_key}')
        # 여기서 필요한 작업을 수행 (예: 다운로드 URL 생성 또는 다른 처리 등)
        return uploaded_file_key
    else:
        print('파일 업로드 실패')