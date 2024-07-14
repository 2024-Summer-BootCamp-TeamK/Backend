import os
from pdf2docx import Converter
from docx import Document
import requests
from dotenv import load_dotenv

load_dotenv()

# S3에서 PDF 파일을 다운로드하여 로컬에 저장하는 함수
def download_pdf_from_s3(url, local_path):
    try:
        response = requests.get(url)
        response.raise_for_status()  # 요청이 성공하지 않을 경우 예외 발생
        with open(local_path, 'wb') as file:
            file.write(response.content)
        print(f"PDF 파일이 {local_path}에 저장되었습니다.")
    except requests.exceptions.RequestException as e:
        print(f"An error occurred: {e}")

# S3 URL 설정
url = f'https://{os.getenv("AWS_STORAGE_BUCKET_NAME")}.s3.ap-northeast-2.amazonaws.com/contracts/5b3bdd00-2f6b-45c5-9ae4-27becc5bad94.pdf'
pdf_file = 'downloaded.pdf'

# PDF 파일을 다운로드하여 로컬에 저장
download_pdf_from_s3(url, pdf_file)

# PDF 파일이 존재하는지 확인
if not os.path.isfile(pdf_file):
    raise FileNotFoundError(f"PDF 파일을 찾을 수 없습니다: {pdf_file}")

# 출력할 Word 파일 경로
docx_file = 'output.docx'

# PDF를 Word로 변환
cv = Converter(pdf_file)
cv.convert(docx_file, start=0, end=None)  # 모든 페이지 변환
cv.close()

# 변환된 Word 파일 열기
doc = Document(docx_file)

# Word 파일의 내용 수정
for paragraph in doc.paragraphs:
    if '고용주는 근로자의 안전과 건강을 보호하기 위해 필요한 조치를 취해야 하며, 근로자는 안전규정을 준수해야 합니다. 안전규정 위반시, 고용주는 즉시 계약을 해지할 수 있습니다. (산업안전보건법 제4조)' in paragraph.text:
        paragraph.text = paragraph.text.replace('고용주는 근로자의 안전과 건강을 보호하기 위해 필요한 조치를 취해야 하며, 근로자는 안전규정을 준수해야 합니다. 안전규정 위반시, 고용주는 즉시 계약을 해지할 수 있습니다. (산업안전보건법 제4조)', '(산업안전보건법 제4조) 고용주는 근로자의 안전을 보장하기 위해 필요한 조치를 취해야 하며, 근로자는 안전규정을 준수해야 합니다. 안전규정 위반 시, 고용주는 즉시 조치를 취해야 합니다.')

# 수정된 Word 파일 저장
modified_docx_file = 'modified_output.docx'
doc.save(modified_docx_file)
print(f"수정된 Word 파일이 {modified_docx_file}에 저장되었습니다.")

# 수정된 Word 파일의 내용 출력
for paragraph in doc.paragraphs:
    print(paragraph.text)
